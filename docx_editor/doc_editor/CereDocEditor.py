"""
CereDocEditor

CereDocEditor is a Python class designed to manipulate and edit .docx documents using the python-docx library.
It provides functionalities such as opening a document, adding statements, deleting text, replacing text,
and saving the document.
"""

import os

from docx import Document
from new_dtect_logging.logger import DtectLogging

import cere_doc_editor.configuration as config

# Logger initialization
log_file_name = (__file__.split(os.sep)[-2] + "_" + __file__.split(os.sep)[-1].split(".")[0])
if __name__ == "__main__":
    lobj_dtect_logger = DtectLogging(log_file_name)
else:
    lobj_dtect_logger = DtectLogging()


class CereDocEditor:
    def __init__(self):
        self.lobj_document = None
        self.lstr_file_path = None

    def open_doc(self, pstr_file_path):
        """Open the .docx file."""
        self.lstr_file_path = pstr_file_path
        try:
            self.lobj_document = Document(self.lstr_file_path)
            print(f"Opened document: {self.lstr_file_path}")
        except Exception as e:
            self._log_warning(f"Error while opening document: {e}")

    def add_statement(self, pstr_statement, pint_page_number_one=-1):
        """Add a statement to the document."""
        if not self.lobj_document:
            print("Document is not opened.")
            return
        try:
            if pint_page_number_one == -1:
                self.lobj_document.add_paragraph(pstr_statement)
                print(f"Added statement: {pstr_statement} to the end of the document")
            else:
                plst_paragraphs = self._get_paragraphs_by_page_number(pint_page_number_one)
                if plst_paragraphs:
                    plst_paragraphs[-1].add_run(f"\n{pstr_statement}")
                    print(f"Added statement: {pstr_statement} to page number: {pint_page_number_one}")
                else:
                    print(f"Page number {pint_page_number_one} does not exist.")
        except Exception as lobj_e:
            self._log_warning(f"Error while adding statement: {lobj_e}")

    def delete_text(self, pstr_text, pint_page_number_two=-1):
        """Delete the first occurrence of the given text from the document."""
        if not self.lobj_document:
            print("Document is not opened.")
            return
        try:
            if pint_page_number_two == -1:
                for pstr_paragraph in self.lobj_document.paragraphs:
                    if pstr_text in pstr_paragraph.text:
                        pstr_paragraph.text = pstr_paragraph.text.replace(pstr_text, "", 1)
                        print(f"Deleted first occurrence of text: {pstr_text}")
                        break
            else:
                llst_paragraphs = self._get_paragraphs_by_page_number(pint_page_number_two)
                if llst_paragraphs:
                    for pstr_paragraph in llst_paragraphs:
                        if pstr_text in pstr_paragraph.text:
                            pstr_paragraph.text = pstr_paragraph.text.replace(pstr_text, "", 1)
                            print(
                                f"Deleted first occurrence of text: {pstr_text} from page number: {pint_page_number_two}")
                            break
                else:
                    print(f"Page number {pint_page_number_two} does not exist.")
        except Exception as e:
            self._log_warning(f"Error while deleting text: {e}")

    def replace_text(self, old_text, new_text):
        """Replace only the first occurrence of old_text with new_text."""
        if not self.lobj_document:
            print("Document is not opened.")
            return
        try:
            for paragraph in self.lobj_document.paragraphs:
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text, 1)
                    # self._add_comment(paragraph, "Replacement done here")
                    print(f"Replaced first occurrence of '{old_text}' with '{new_text}'")
                    break  # Stop after the first replacement
        except Exception as e:
            self._log_warning(f"Error while replacing text: {e}")

    def add_text_after_match(self, match_text, new_text):
        """Add new_text after the first occurrence of match_text."""
        if not self.lobj_document:
            print("Document is not opened.")
            return
        try:
            for paragraph in self.lobj_document.paragraphs:
                if match_text in paragraph.text:
                    index = paragraph.text.find(match_text) + len(match_text)
                    paragraph.text = paragraph.text[:index] + new_text + paragraph.text[index:]
                    print(f"Added '{new_text}' after the first occurrence of '{match_text}'")
                    break  # Stop after the first addition
        except Exception as e:
            self._log_warning(f"Error while adding text after match: {e}")

    def save_doc(self):
        """Save the document."""
        if not self.lobj_document:
            print("Document is not opened.")
            return
        try:
            self.lobj_document.save(self.lstr_file_path)
            print(f"Document saved: {self.lstr_file_path}")
        except Exception as e:
            self._log_warning(f"Error while saving document: {e}")

    # def _get_paragraphs_by_page_number(self, pint_page_number_three):
    #     """Helper method to get paragraphs by page number."""
    #     try:
    #         llst_paragraphs = [lstr_p for lstr_p in self.lobj_document.paragraphs if
    #                            lstr_p.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY]
    #         return [llst_paragraphs[pint_page_number_three - 1]] if pint_page_number_three - 1 < len(
    #             llst_paragraphs) else []
    #     except Exception as e:
    #         self._log_warning(f"Error while retrieving paragraphs: {e}")
    #         return []

    def _log_warning(self, pstr_message):
        """Log warnings with standardized format."""
        lobj_dtect_logger.logger.warning(pstr_message)
        print(pstr_message)

    # def _add_comment(self, paragraph, comment_text):
    #     """Add a comment to the specified paragraph using core properties."""
    #     core_properties = self.document.core_properties
    #     core_properties.comments = comment_text
    #     core_properties.last_modified_by = "Auto Commenter"
    #     core_properties.modified = datetime.now()


if __name__ == "__main__":
    file_path = config.INPUT_PATH

    editor = CereDocEditor()
    editor.open_doc(file_path)

    # # Example usage
    editor.replace_text("Model", "SAMPLE REPLACEMENT")
    editor.add_text_after_match("Overview", " AFTER TEXT After overview")
    editor.delete_text("#hihi - CERE SAMPLE#")

    editor.save_doc()
